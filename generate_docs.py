from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return h

def para(text="", bold_prefix=None, code=False):
    """Add a paragraph, optionally with a bold prefix label."""
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    r2 = p.add_run(text)
    if code:
        r2.font.name = "Courier New"
        r2.font.size = Pt(9)
    return p

def bullet(text, sub=False):
    style = "List Bullet 2" if sub else "List Bullet"
    try:
        p = doc.add_paragraph(text, style=style)
    except KeyError:
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.5 if sub else 0.25)
    return p

def divider():
    doc.add_paragraph("─" * 60)

# ─── Title ───────────────────────────────────────────────────────────────────

title = doc.add_heading("Portfolio Project Docs", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("A concise reference for the tech used, how to set up, and key implementation details.").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("")

# ─── 1. Tech Stack Overview ──────────────────────────────────────────────────

set_heading("1. Tech Stack & Significance", 1)

stack = [
    ("Next.js 16",         "React framework — App Router, SSR, file-based routing, API routes."),
    ("React 19",           "UI library; concurrent features, hooks."),
    ("TypeScript 5",       "Static typing across the whole codebase."),
    ("Tailwind CSS 4",     "Utility-first CSS; all styling done inline via class names."),
    ("Framer Motion 12",   "Animation library — fade-ins, scroll-triggered reveals, count-up numbers."),
    ("React Hook Form 7",  "Performant form state management with minimal re-renders."),
    ("Zod 4",              "Schema validation; shared between client and API route."),
    ("@hookform/resolvers","Bridges Zod schema into React Hook Form validation."),
    ("Resend 6",           "Email API — sends contact form submissions server-side."),
    ("clsx + tailwind-merge","Conditionally combine class names without duplication conflicts."),
    ("lucide-react",       "Icon library (SVG-based, tree-shakable)."),
]

for name, desc in stack:
    p = doc.add_paragraph()
    p.add_run(f"{name}  ").bold = True
    p.add_run(desc)

doc.add_paragraph("")

# ─── 2. Installation ─────────────────────────────────────────────────────────

set_heading("2. Installation & Scripts", 1)

para("Clone & install dependencies:", bold_prefix="Step 1  ")
p = doc.add_paragraph()
r = p.add_run(
    "git clone <repo-url>\n"
    "cd portfolio\n"
    "npm install"
)
r.font.name = "Courier New"
r.font.size = Pt(9)

para("\nCreate a .env.local file in the project root:", bold_prefix="Step 2  ")
p = doc.add_paragraph()
r = p.add_run("RESEND_API_KEY=your_resend_key\nCONTACT_TO_EMAIL=you@example.com")
r.font.name = "Courier New"
r.font.size = Pt(9)

para("\nRun the dev server:", bold_prefix="Step 3  ")
p = doc.add_paragraph()
r = p.add_run("npm run dev      # http://localhost:3000\nnpm run build    # production build\nnpm start        # serve production build")
r.font.name = "Courier New"
r.font.size = Pt(9)

doc.add_paragraph("")

# ─── 3. Project Structure ────────────────────────────────────────────────────

set_heading("3. Project Structure (key files)", 1)

structure = [
    ("src/app/page.tsx",                 "Root page — assembles all sections."),
    ("src/app/layout.tsx",               "Global HTML shell, fonts, metadata."),
    ("src/app/globals.css",              "Tailwind base + custom CSS variables."),
    ("src/app/api/contact/route.ts",     "API route for the contact form (Resend)."),
    ("src/data/portfolio.ts",            "Single source of truth — all content (projects, skills, links)."),
    ("src/components/sections/",         "Hero, About, Experience, Projects, TechStack, Contact."),
    ("src/components/animations.tsx",    "Reusable Framer Motion wrappers: FadeIn, StaggerContainer, CountUp."),
    ("src/components/SkillCarousel.tsx", "CSS 3D rotating skill carousel."),
    ("src/components/CommandPalette.tsx","Terminal-style search/navigation modal."),
    ("src/components/CommandPaletteProvider.tsx","Mounts the palette, listens for Cmd/Ctrl+K."),
    ("src/hooks/useContactForm.ts",      "Form logic hook — validation, submission, state."),
    ("src/lib/contactSchema.ts",         "Zod schema for the contact form."),
    ("src/lib/utils.ts",                 "cn() helper — merges Tailwind classes safely."),
]

for path, desc in structure:
    p = doc.add_paragraph()
    p.add_run(f"{path}  ").bold = True
    r = p.add_run(f"— {desc}")

doc.add_paragraph("")

# ─── 4. Skill Carousel ───────────────────────────────────────────────────────

set_heading("4. Skill Carousel — How It Works", 1)

doc.add_paragraph(
    "The carousel is a pure CSS 3D cylinder — no external library. "
    "Each skill card is placed on the surface of an imaginary cylinder using CSS 3D transforms, "
    "then the whole track rotates with a CSS keyframe animation."
)

set_heading("4.1 Core Math", 2)
doc.add_paragraph("Given N cards and a desired card spacing of ~145 px:")
p = doc.add_paragraph()
r = p.add_run(
    "RADIUS = round( (145 × N) / (2π) )\n"
    "PERSPECTIVE = RADIUS × 4.5\n"
    "DURATION = N × 1.8 seconds  (≈ comfortable spin speed)"
)
r.font.name = "Courier New"
r.font.size = Pt(9)

set_heading("4.2 Card Placement", 2)
doc.add_paragraph("Each card i gets a unique Y-axis rotation and is pushed out by RADIUS:")
p = doc.add_paragraph()
r = p.add_run(
    "angle = (360 / N) × i\n"
    "transform: rotateY(${angle}deg) translateZ(${RADIUS}px)"
)
r.font.name = "Courier New"
r.font.size = Pt(9)
doc.add_paragraph(
    "Because transform-style: preserve-3d is set on the parent track, each card "
    "inherits the 3D context and appears at its correct position on the cylinder."
)

set_heading("4.3 Spin Animation", 2)
p = doc.add_paragraph()
r = p.add_run(
    "@keyframes carousel-spin {\n"
    "  from { transform: rotateY(0deg); }\n"
    "  to   { transform: rotateY(-360deg); }\n"
    "}\n"
    ".carousel-track {\n"
    "  animation: carousel-spin <DURATION>s linear infinite;\n"
    "}\n"
    ".carousel-track:hover {\n"
    "  animation-play-state: paused;\n"
    "}"
)
r.font.name = "Courier New"
r.font.size = Pt(9)
doc.add_paragraph("Hover pauses the animation via animation-play-state: paused — zero JS needed.")

set_heading("4.4 Edge Fading", 2)
doc.add_paragraph(
    "Two absolutely-positioned divs with gradient backgrounds (from-[#111111] to-transparent) "
    "are layered on left and right with pointer-events: none so they fade cards in/out "
    "without blocking interaction."
)
doc.add_paragraph("")

# ─── 5. Command Palette ──────────────────────────────────────────────────────

set_heading("5. Command Palette — How It Works", 1)

doc.add_paragraph(
    "A terminal-themed search modal (Cmd/Ctrl+K) that lets visitors navigate sections, "
    "open projects on GitHub, and run pseudo shell commands."
)

set_heading("5.1 Architecture — Two Components", 2)

p = doc.add_paragraph()
p.add_run("CommandPaletteProvider  ").bold = True
p.add_run("— Thin wrapper. Holds isOpen state, listens for the keyboard shortcut, "
          "and passes an openCount key to CommandPalette so its internal state fully resets on every open.")

p = doc.add_paragraph()
p.add_run("CommandPalette  ").bold = True
p.add_run("— The actual UI. Receives isOpen and onClose props; renders nothing when closed (early return null).")

set_heading("5.2 Data & Items", 2)
doc.add_paragraph("Items are built once with useMemo and never re-computed unless onClose changes:")
bullet("Sections — nav links → smooth-scroll via scrollIntoView")
bullet("Projects — portfolio data → opens githubUrl in new tab")
bullet("Commands — ls (list all) and clear (wipe history)")

set_heading("5.3 Search / Filtering Logic", 2)
p = doc.add_paragraph()
r = p.add_run(
    "matchesSearch(item, query):\n"
    "  haystack = [label, command, ...keywords].join(' ').toLowerCase()\n"
    "  every word in query must appear in haystack (AND match)"
)
r.font.name = "Courier New"
r.font.size = Pt(9)
doc.add_paragraph("Prefix shortcuts are also supported:")
bullet('cd <term>  — filters to Sections only')
bullet('open <term>  — filters to Projects only')
bullet('ls  — shows all non-command items')
bullet('clear  — clears terminal history')

set_heading("5.4 Keyboard Navigation", 2)
pairs = [
    ("↑ / ↓",    "Move selected index up/down through filtered list."),
    ("Tab",       "Auto-complete the selected item's command into the input."),
    ("Enter",     "Run the selected (or first) item's action."),
    ("Escape",    "Close the palette."),
]
for key, desc in pairs:
    p = doc.add_paragraph()
    p.add_run(f"{key}  ").bold = True
    p.add_run(desc)

set_heading("5.5 Terminal History", 2)
doc.add_paragraph(
    "Every executed command is appended to a history array (useState). "
    "The history is rendered above the input just like a real terminal output log. "
    "The 'clear' command empties this array."
)

set_heading("5.6 UI Details", 2)
bullet("macOS-style traffic-light dots in the titlebar (cosmetic only).")
bullet("Backdrop blur + black overlay; clicking outside calls onClose.")
bullet("Items grouped by Sections / Projects / Commands with colour-coded labels.")
bullet("Scrollable results list (max-h + overflow-y-auto).")

doc.add_paragraph("")
divider()
doc.add_paragraph("End of document — generated for portfolio@0.1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── Save ────────────────────────────────────────────────────────────────────

doc.save("/Users/sudip/PROJECTS/personal-project/portfolio/project-docs.docx")
print("Done → project-docs.docx")
